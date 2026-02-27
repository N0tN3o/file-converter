import sys
import os
import zipfile
import filetype
import fitz
import docx
import markdown
from xhtml2pdf import pisa
from PIL import Image
import qtawesome as qta

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QFileDialog, QComboBox, QProgressBar, QMessageBox,
    QGroupBox, QFrame
)
from PySide6.QtCore import QThread, Signal, Qt, QSize
from PySide6.QtGui import QFont, QIcon, QStyleHints

# --- Theme Stylesheets ---
DARK_STYLESHEET = """
QMainWindow {
    background-color: #1e1e2e;
}
QLabel {
    color: #cdd6f4;
    font-size: 13px;
}
QGroupBox {
    color: #bac2de;
    border: 1px solid #313244;
    border-radius: 8px;
    margin-top: 12px;
    font-weight: bold;
    padding-top: 15px;
}
QGroupBox::title {
    subcontrol-origin: margin;
    subcontrol-position: top left;
    left: 15px;
    padding: 0 5px;
}
QPushButton {
    background-color: #313244;
    color: #cdd6f4;
    border-radius: 6px;
    padding: 8px 16px;
    font-weight: bold;
    border: none;
}
QPushButton:hover {
    background-color: #45475a;
}
QPushButton#convertBtn {
    background-color: #89b4fa;
    color: #11111b;
    font-size: 15px;
    padding: 12px;
    border-radius: 8px;
}
QPushButton#convertBtn:hover {
    background-color: #b4befe;
}
QPushButton#convertBtn:disabled {
    background-color: #313244;
    color: #585b70;
}
QComboBox {
    background-color: #313244;
    color: #cdd6f4;
    border-radius: 4px;
    padding: 6px 10px;
    border: 1px solid #45475a;
}
QComboBox:disabled {
    background-color: #181825;
    color: #45475a;
    border: 1px solid #313244;
}
QComboBox::drop-down {
    border: none;
}
QProgressBar {
    border: 1px solid #313244;
    border-radius: 6px;
    text-align: center;
    color: #cdd6f4;
    background-color: #181825;
    font-weight: bold;
}
QProgressBar::chunk {
    background-color: #a6e3a1;
    border-radius: 5px;
}
#dropZone {
    border: 2px dashed #45475a;
    border-radius: 10px;
    background-color: #181825;
}
"""

LIGHT_STYLESHEET = """
QMainWindow {
    background-color: #eff1f5;
}
QLabel {
    color: #4c4f69;
    font-size: 13px;
}
QGroupBox {
    color: #5c5f77;
    border: 1px solid #ccd0da;
    border-radius: 8px;
    margin-top: 12px;
    font-weight: bold;
    padding-top: 15px;
}
QGroupBox::title {
    subcontrol-origin: margin;
    subcontrol-position: top left;
    left: 15px;
    padding: 0 5px;
}
QPushButton {
    background-color: #ccd0da;
    color: #4c4f69;
    border-radius: 6px;
    padding: 8px 16px;
    font-weight: bold;
    border: none;
}
QPushButton:hover {
    background-color: #bcc0cc;
}
QPushButton#convertBtn {
    background-color: #1e66f5;
    color: #eff1f5;
    font-size: 15px;
    padding: 12px;
    border-radius: 8px;
}
QPushButton#convertBtn:hover {
    background-color: #7287fd;
}
QPushButton#convertBtn:disabled {
    background-color: #ccd0da;
    color: #9ca0b0;
}
QComboBox {
    background-color: #ccd0da;
    color: #4c4f69;
    border-radius: 4px;
    padding: 6px 10px;
    border: 1px solid #bcc0cc;
}
QComboBox:disabled {
    background-color: #e6e9ef;
    color: #9ca0b0;
    border: 1px solid #ccd0da;
}
QComboBox::drop-down {
    border: none;
}
QProgressBar {
    border: 1px solid #ccd0da;
    border-radius: 6px;
    text-align: center;
    color: #4c4f69;
    background-color: #e6e9ef;
    font-weight: bold;
}
QProgressBar::chunk {
    background-color: #40a02b;
    border-radius: 5px;
}
#dropZone {
    border: 2px dashed #bcc0cc;
    border-radius: 10px;
    background-color: #e6e9ef;
}
"""

# Icon color palettes per theme
THEME_COLORS = {
    "dark": {
        "text": "#cdd6f4",
        "subtext": "#a6adc8",
        "overlay": "#6c7086",
        "muted": "#585b70",
        "green": "#a6e3a1",
        "blue": "#89b4fa",
        "surface": "#181825",
        "drop_bg": "#1e1e2e",
        "instructions": "#cdd6f4",
    },
    "light": {
        "text": "#4c4f69",
        "subtext": "#6c6f85",
        "overlay": "#9ca0b0",
        "muted": "#9ca0b0",
        "green": "#40a02b",
        "blue": "#1e66f5",
        "surface": "#e6e9ef",
        "drop_bg": "#dce0e8",
        "instructions": "#4c4f69",
    },
}

CODE_EXTS = ["py", "js", "c", "cpp", "cs", "java", "json", "css", "html"]

CONVERSIONS = {
    "PDF": ["png", "jpg", "jpeg", "txt"],
    "Image": ["png", "jpg", "jpeg", "webp", "bmp", "gif"],
    "DOCX": ["txt", "md"],
    "TXT": ["docx", "md"] + CODE_EXTS,
    "MD": ["txt", "docx", "html", "pdf"],
    "HTML": ["pdf", "txt", "md"],
    "Code": ["txt", "md", "docx"] 
}

class ConvertWorker(QThread):
    progress = Signal(int)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, input_file, output_dir, source_format, target_format):
        super().__init__()
        self.input_file = input_file
        self.output_dir = output_dir
        self.source_format = source_format
        self.target_format = target_format.lower()

    def run(self):
        try:
            base_name = os.path.splitext(os.path.basename(self.input_file))[0]

            if self.source_format == "PDF":
                doc = fitz.open(self.input_file)
                total_pages = len(doc)

                if self.target_format == "txt":
                    text_content = ""
                    for i in range(total_pages):
                        text_content += doc.load_page(i).get_text()
                        self.progress.emit(int(((i + 1) / total_pages) * 100))
                    
                    out_path = os.path.join(self.output_dir, f"{base_name}.txt")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(text_content)
                    self.finished.emit("PDF successfully extracted to Text!")

                elif self.target_format in ["png", "jpg", "jpeg"]:
                    if total_pages > 1:
                        zip_path = os.path.join(self.output_dir, f"{base_name}_images.zip")
                        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                            for i in range(total_pages):
                                page = doc.load_page(i)
                                pix = page.get_pixmap(dpi=300)
                                img_data = pix.tobytes(self.target_format)
                                zipf.writestr(f"{base_name}_page_{i+1}.{self.target_format}", img_data)
                                self.progress.emit(int(((i + 1) / total_pages) * 100))
                        self.finished.emit("Multiple PDF pages successfully zipped as images!")
                    else:
                        page = doc.load_page(0)
                        pix = page.get_pixmap(dpi=300)
                        out_path = os.path.join(self.output_dir, f"{base_name}.{self.target_format}")
                        pix.save(out_path)
                        self.progress.emit(100)
                        self.finished.emit("Single PDF page successfully converted to image!")
                doc.close()

            elif self.source_format == "Image":
                with Image.open(self.input_file) as img:
                    if self.target_format in ['jpg', 'jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
                        img = img.convert('RGB')
                    out_path = os.path.join(self.output_dir, f"{base_name}_converted.{self.target_format}")
                    img.save(out_path)
                    self.progress.emit(100)
                    self.finished.emit(f"Image converted to {self.target_format.upper()}!")

            elif self.source_format == "DOCX":
                doc = docx.Document(self.input_file)
                text_content = "\n".join([p.text for p in doc.paragraphs])
                
                out_path = os.path.join(self.output_dir, f"{base_name}.{self.target_format}")
                with open(out_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                self.progress.emit(100)
                self.finished.emit(f"DOCX converted to {self.target_format.upper()}!")

            elif self.source_format in ["TXT", "MD", "HTML", "Code"]:
                with open(self.input_file, 'r', encoding='utf-8') as f:
                    text_content = f.read()

                if self.target_format == "docx":
                    doc = docx.Document()
                    doc.add_paragraph(text_content)
                    out_path = os.path.join(self.output_dir, f"{base_name}.docx")
                    doc.save(out_path)
                
                elif self.target_format == "html" and self.source_format == "MD":
                    html_content = markdown.markdown(text_content)
                    out_path = os.path.join(self.output_dir, f"{base_name}.html")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(html_content)

                elif self.target_format == "pdf" and self.source_format in ["MD", "HTML"]:
                    if self.source_format == "MD":
                        html_content = markdown.markdown(text_content)
                    else:
                        html_content = text_content
                        
                    out_path = os.path.join(self.output_dir, f"{base_name}.pdf")
                    with open(out_path, "wb") as pdf_file:
                        pisa.CreatePDF(html_content, dest=pdf_file)

                else:
                    out_path = os.path.join(self.output_dir, f"{base_name}.{self.target_format}")
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(text_content)
                
                self.progress.emit(100)
                self.finished.emit(f"File converted to {self.target_format.upper()}!")

        except Exception as e:
            self.error.emit(str(e))


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AnyDoc - Universal File Converter")
        self.resize(550, 480)
        self.input_file = None
        self.output_dir = None
        self.custom_output_dir = False
        self.is_dark = True
        self.setAcceptDrops(True)
        
        self.init_ui()
        self._detect_and_apply_theme()
        self._connect_theme_signal()

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(15)
        main_layout.setContentsMargins(20, 20, 20, 20)

        # --- Instructions (Replaced Top Header) ---
        self.lbl_instructions = QLabel("Select a file to convert, verify the format settings, and click Start Conversion.")
        self.lbl_instructions.setWordWrap(True)
        main_layout.addWidget(self.lbl_instructions)

        # --- File Selection Zone ---
        self.file_frame = QFrame()
        self.file_frame.setObjectName("dropZone")
        file_layout = QVBoxLayout(self.file_frame)
        file_layout.setContentsMargins(15, 20, 15, 20)
        
        self.lbl_file_icon = QLabel()
        self.lbl_file_icon.setAlignment(Qt.AlignCenter)
        file_layout.addWidget(self.lbl_file_icon)

        self.lbl_file = QLabel("No file selected")
        self.lbl_file.setAlignment(Qt.AlignCenter)
        file_layout.addWidget(self.lbl_file)

        self.lbl_drop_hint = QLabel("or drag && drop a file here")
        self.lbl_drop_hint.setAlignment(Qt.AlignCenter)
        file_layout.addWidget(self.lbl_drop_hint)

        self.btn_select_file = QPushButton(" Browse Files")
        self.btn_select_file.setCursor(Qt.PointingHandCursor)
        self.btn_select_file.clicked.connect(self.select_file)
        
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        btn_layout.addWidget(self.btn_select_file)
        btn_layout.addStretch()
        file_layout.addLayout(btn_layout)
        
        main_layout.addWidget(self.file_frame)

        # --- Settings Group ---
        settings_group = QGroupBox("Conversion Settings")
        settings_layout = QVBoxLayout(settings_group)
        settings_layout.setSpacing(10)

        # Formats
        format_layout = QHBoxLayout()
        
        self.combo_source = QComboBox()
        self.combo_source.addItems(list(CONVERSIONS.keys()))
        self.combo_source.setEnabled(False)
        self.combo_source.currentTextChanged.connect(self.update_target_formats)
        
        self.combo_target = QComboBox()
        self.combo_target.setEnabled(False)

        format_layout.addWidget(QLabel("Source:"))
        format_layout.addWidget(self.combo_source, 1)
        
        self.arrow_lbl = QLabel()
        format_layout.addWidget(self.arrow_lbl)
        
        format_layout.addWidget(QLabel("Target:"))
        format_layout.addWidget(self.combo_target, 1)
        
        settings_layout.addLayout(format_layout)

        # Output Dir
        dir_layout = QHBoxLayout()
        self.lbl_dir = QLabel("Output: Same as source folder")
        self.lbl_dir.setStyleSheet("color: #a6adc8; font-style: italic;")
        
        self.btn_select_dir = QPushButton()
        self.btn_select_dir.setToolTip("Select Custom Output Folder")
        self.btn_select_dir.setCursor(Qt.PointingHandCursor)
        self.btn_select_dir.clicked.connect(self.select_output_dir)
        
        dir_layout.addWidget(self.lbl_dir, 1)
        dir_layout.addWidget(self.btn_select_dir)
        settings_layout.addLayout(dir_layout)

        main_layout.addWidget(settings_group)

        # --- Bottom Area ---
        main_layout.addStretch()

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setFixedHeight(10)
        main_layout.addWidget(self.progress_bar)

        self.btn_convert = QPushButton("Start Conversion")
        self.btn_convert.setObjectName("convertBtn")
        self.btn_convert.setEnabled(False)
        self.btn_convert.setCursor(Qt.PointingHandCursor)
        self.btn_convert.clicked.connect(self.start_conversion)
        main_layout.addWidget(self.btn_convert)


    # --- Theme Management ---
    def _detect_and_apply_theme(self):
        """Detect the OS color scheme and apply the matching theme."""
        try:
            scheme = QApplication.instance().styleHints().colorScheme()
            self.is_dark = (scheme == Qt.ColorScheme.Dark)
        except AttributeError:
            # Qt < 6.5 fallback: assume dark
            self.is_dark = True
        self._apply_theme()

    def _connect_theme_signal(self):
        """Connect to the OS theme-change signal for live switching."""
        try:
            QApplication.instance().styleHints().colorSchemeChanged.connect(
                self._on_system_theme_changed
            )
        except AttributeError:
            pass  # Qt < 6.5, no reactive support

    def _on_system_theme_changed(self, scheme):
        self.is_dark = (scheme == Qt.ColorScheme.Dark)
        self._apply_theme()

    def _apply_theme(self):
        """Apply current theme stylesheet and refresh all themed icons."""
        c = THEME_COLORS["dark" if self.is_dark else "light"]
        QApplication.instance().setStyleSheet(
            DARK_STYLESHEET if self.is_dark else LIGHT_STYLESHEET
        )

        # Instruction label
        self.lbl_instructions.setStyleSheet(
            f"font-size: 14px; color: {c['instructions']}; margin-bottom: 5px;"
        )

        # File icon (show check if a file is loaded, otherwise default)
        if self.input_file:
            self.lbl_file_icon.setPixmap(
                qta.icon('fa5s.check-circle', color=c['green']).pixmap(QSize(40, 40))
            )
        else:
            self.lbl_file_icon.setPixmap(
                qta.icon('fa5s.file-alt', color=c['overlay']).pixmap(QSize(40, 40))
            )

        # Text labels
        self.lbl_file.setStyleSheet(f"color: {c['subtext']};")
        self.lbl_drop_hint.setStyleSheet(
            f"color: {c['muted']}; font-size: 11px; font-style: italic;"
        )
        self.lbl_dir.setStyleSheet(f"color: {c['subtext']}; font-style: italic;")

        # Buttons / icons
        self.btn_select_file.setIcon(qta.icon('fa5s.folder-open', color=c['text']))
        self.btn_select_dir.setIcon(qta.icon('fa5s.folder', color=c['text']))
        self.arrow_lbl.setPixmap(
            qta.icon('fa5s.arrow-right', color=c['overlay']).pixmap(QSize(16, 16))
        )

    # --- Drag & Drop Events ---
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            c = THEME_COLORS["dark" if self.is_dark else "light"]
            self.file_frame.setStyleSheet(
                f"#dropZone {{ border: 2px dashed {c['blue']}; border-radius: 10px; "
                f"background-color: {c['drop_bg']}; }}"
            )

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dragLeaveEvent(self, event):
        self.file_frame.setStyleSheet("")

    def dropEvent(self, event):
        self.file_frame.setStyleSheet("")
        urls = event.mimeData().urls()
        if urls:
            file_path = urls[0].toLocalFile()
            if os.path.isfile(file_path):
                self.load_file(file_path)

    # --- File Handling ---
    def select_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select File to Convert")
        if file_path:
            self.load_file(file_path)

    def load_file(self, file_path):
        """Shared logic for both browsing and drag-and-drop"""
        self.input_file = file_path
        self.lbl_file.setText(f"<b>{os.path.basename(file_path)}</b>")
        c = THEME_COLORS["dark" if self.is_dark else "light"]
        self.lbl_file_icon.setPixmap(qta.icon('fa5s.check-circle', color=c['green']).pixmap(QSize(40, 40)))
        
        if not self.custom_output_dir:
            self.output_dir = os.path.dirname(file_path)

        self.detect_file_type()

    def select_output_dir(self):
        dir_path = QFileDialog.getExistingDirectory(self, "Select Output Folder")
        if dir_path:
            self.output_dir = dir_path
            self.custom_output_dir = True 
            
            # Truncate long paths for visual cleanliness
            display_path = dir_path if len(dir_path) < 40 else "..." + dir_path[-37:]
            self.lbl_dir.setText(f"Output: {display_path}")

    def detect_file_type(self):
        self.combo_source.setEnabled(True)
        self.lbl_drop_hint.setVisible(False)
        kind = filetype.guess(self.input_file)
        ext = os.path.splitext(self.input_file)[1].lower()

        detected_type = None
        if kind and "pdf" in kind.mime:
            detected_type = "PDF"
        elif kind and "image" in kind.mime:
            detected_type = "Image"
        elif ext == ".pdf":
            detected_type = "PDF"
        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"]:
            detected_type = "Image"
        elif ext == ".docx":
            detected_type = "DOCX"
        elif ext == ".md":
            detected_type = "MD"
        elif ext in [".htm", ".html"]:
            detected_type = "HTML"
        elif ext == ".txt":
            detected_type = "TXT"
        elif ext.replace(".", "") in CODE_EXTS:
            detected_type = "Code"

        if detected_type:
            index = self.combo_source.findText(detected_type)
            if index >= 0:
                self.combo_source.setCurrentIndex(index)
        else:
            QMessageBox.warning(self, "Unknown File", "Could not reliably detect file type. Please manually select the Source Format.")
            self.combo_source.setCurrentIndex(0)
    def update_target_formats(self, source_format):
        self.combo_target.clear()
        if source_format in CONVERSIONS:
            self.combo_target.addItems(CONVERSIONS[source_format])
            self.combo_target.setEnabled(True)
            self.btn_convert.setEnabled(True)
        else:
            self.combo_target.setEnabled(False)
            self.btn_convert.setEnabled(False)

    def start_conversion(self):
        source_fmt = self.combo_source.currentText()
        target_fmt = self.combo_target.currentText()
        
        self.btn_convert.setEnabled(False)
        self.btn_convert.setText("  Converting...")
        self.progress_bar.setValue(0)

        self.worker = ConvertWorker(self.input_file, self.output_dir, source_fmt, target_fmt)
        self.worker.progress.connect(self.update_progress)
        self.worker.finished.connect(self.conversion_finished)
        self.worker.error.connect(self.conversion_error)
        self.worker.start()

    def update_progress(self, val):
        self.progress_bar.setValue(val)

    def conversion_finished(self, msg):
        self.btn_convert.setEnabled(True)
        self.btn_convert.setText("  Start Conversion")
        self.progress_bar.setValue(100)
        QMessageBox.information(self, "Success", msg)
        self.progress_bar.setValue(0)

    def conversion_error(self, err_msg):
        self.btn_convert.setEnabled(True)
        self.btn_convert.setText("  Start Conversion")
        QMessageBox.critical(self, "Error", f"An error occurred:\n{err_msg}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    
    # Stylesheet is applied by MainWindow._apply_theme()
    
    # Set global default font
    font = QFont("Segoe UI", 10)
    app.setFont(font)
    
    window = MainWindow()
    window.show()
    sys.exit(app.exec())